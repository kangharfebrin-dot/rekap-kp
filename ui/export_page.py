"""
ui/export_page.py
ExportMixin — berisi UI dan logika untuk halaman export (Excel & Word).
"""

import os
import customtkinter as ctk
from tkinter import messagebox, filedialog


class ExportMixin:

    def build_export_page(self):
        page = ctk.CTkScrollableFrame(self.main_frame, fg_color="transparent")
        self.pages["export"] = page

        ctk.CTkLabel(page, text="Export Laporan",
                     font=ctk.CTkFont(size=20, weight="bold")).pack(anchor="w", pady=(0, 15))

        # Ringkasan data
        stats = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        stats.pack(fill="x", pady=8)
        ctk.CTkLabel(stats, text="📊 Ringkasan Data",
                     font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e94560").pack(anchor="w", padx=15, pady=(12, 5))
        self.stats_label = ctk.CTkLabel(stats, text="", justify="left",
                                        font=ctk.CTkFont(size=13))
        self.stats_label.pack(anchor="w", padx=15, pady=(0, 12))

        # Tombol export
        exp = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        exp.pack(fill="x", pady=8)
        ctk.CTkLabel(exp, text="📤 Pilih Format Export",
                     font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e94560").pack(anchor="w", padx=15, pady=(12, 5))

        for txt, color, hover, cmd in [
            ("📊  Export ke Excel (.xlsx)", "#27ae60", "#1e8449", self.export_excel),
            ("📄  Export ke Word (.docx)", "#2980b9", "#1f6fa3", self.export_word)
        ]:
            ctk.CTkButton(exp, text=txt, font=ctk.CTkFont(size=14, weight="bold"),
                          height=50, fg_color=color, hover_color=hover,
                          command=cmd).pack(fill="x", padx=15, pady=6)

        ctk.CTkLabel(exp, text="", height=10).pack()  # spacer

        # Update stats saat halaman ditampilkan
        page.bind("<Map>", lambda e: self.update_stats())

    def update_stats(self):
        if not self.data:
            return
        p = self.data["profil"]
        n = len(self.data["kegiatan"])
        weeks = set(k["minggu"] for k in self.data["kegiatan"])
        txt = (f"Nama: {p['nama'] or '-'}  |  NIM: {p['nim'] or '-'}\n"
               f"Perusahaan: {p['perusahaan'] or '-'}\n"
               f"Total Kegiatan: {n}  |  Jumlah Minggu: {len(weeks)}")
        self.stats_label.configure(text=txt)

    def export_excel(self):
        if not self.data["kegiatan"]:
            messagebox.showwarning("Peringatan", "Belum ada data kegiatan!")
            return
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile="Rekap_KP.xlsx"
        )
        if not filepath:
            return
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "Rekap Kegiatan KP"

            p = self.data["profil"]
            ws.merge_cells("A1:E1")
            ws["A1"] = "REKAP KEGIATAN KERJA PRAKTEK"
            ws["A1"].font = Font(size=14, bold=True)
            ws["A1"].alignment = Alignment(horizontal="center")

            info = [
                f"Nama: {p['nama']}", f"NIM: {p['nim']}", f"Prodi: {p['prodi']}",
                f"Perusahaan: {p['perusahaan']}",
                f"Periode: {p['periode_mulai']} - {p['periode_selesai']}"
            ]
            for i, txt in enumerate(info):
                ws.merge_cells(f"A{i + 2}:E{i + 2}")
                ws[f"A{i + 2}"] = txt

            start_row = len(info) + 3
            headers = ["No", "Tanggal", "Minggu", "Kegiatan", "Keterangan"]
            header_fill = PatternFill(start_color="E94560", end_color="E94560", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            thin = Side(style="thin")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            for j, h in enumerate(headers):
                cell = ws.cell(row=start_row, column=j + 1, value=h)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = border

            for i, k in enumerate(self.data["kegiatan"]):
                row = start_row + 1 + i
                vals = [i + 1, k["tanggal"], f"Minggu {k['minggu']}", k["kegiatan"], k["keterangan"]]
                for j, v in enumerate(vals):
                    cell = ws.cell(row=row, column=j + 1, value=v)
                    cell.border = border
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            ws.column_dimensions["A"].width = 5
            ws.column_dimensions["B"].width = 14
            ws.column_dimensions["C"].width = 12
            ws.column_dimensions["D"].width = 45
            ws.column_dimensions["E"].width = 30

            wb.save(filepath)
            if messagebox.askyesno("Sukses",
                                   f"File Excel berhasil disimpan!\n{filepath}\n\n"
                                   f"Apakah Anda ingin membuka file tersebut sekarang?"):
                try:
                    os.startfile(filepath)
                except Exception as e:
                    messagebox.showerror("Error", f"Gagal membuka file:\n{e}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def export_word(self):
        if not self.data["kegiatan"]:
            messagebox.showwarning("Peringatan", "Belum ada data kegiatan!")
            return
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialfile="Laporan_KP.docx"
        )
        if not filepath:
            return
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            p = self.data["profil"]

            title = doc.add_heading("LAPORAN KEGIATAN KERJA PRAKTEK", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading("Data Mahasiswa", level=2)
            for label, key in [("Nama", "nama"), ("NIM", "nim"), ("Program Studi", "prodi"),
                                ("Fakultas", "fakultas"), ("Universitas", "universitas")]:
                doc.add_paragraph(f"{label}: {p.get(key, '-')}")

            doc.add_heading("Data Tempat Kerja Praktek", level=2)
            for label, key in [("Perusahaan/Instansi", "perusahaan"),
                                ("Alamat", "alamat_perusahaan"),
                                ("Periode", None),
                                ("Dosen Pembimbing", "dosen_pembimbing"),
                                ("Pembimbing Lapangan", "pembimbing_lapangan")]:
                if key:
                    doc.add_paragraph(f"{label}: {p.get(key, '-')}")
                else:
                    doc.add_paragraph(
                        f"Periode: {p.get('periode_mulai', '-')} s/d {p.get('periode_selesai', '-')}"
                    )

            doc.add_heading("Catatan Kegiatan Harian", level=2)
            weeks = {}
            for k in self.data["kegiatan"]:
                w = k["minggu"]
                weeks.setdefault(w, []).append(k)

            for w in sorted(weeks.keys(), key=lambda x: int(x)):
                doc.add_heading(f"Minggu ke-{w}", level=3)
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, txt in enumerate(["No", "Tanggal", "Kegiatan", "Keterangan"]):
                    cell = table.rows[0].cells[i]
                    cell.text = txt
                    for par in cell.paragraphs:
                        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in par.runs:
                            run.bold = True

                for j, k in enumerate(weeks[w]):
                    row = table.add_row()
                    row.cells[0].text = str(j + 1)
                    row.cells[1].text = k["tanggal"]
                    row.cells[2].text = k["kegiatan"]
                    row.cells[3].text = k["keterangan"]

                doc.add_paragraph("")  # spacer

            doc.save(filepath)
            if messagebox.askyesno("Sukses",
                                   f"File Word berhasil disimpan!\n{filepath}\n\n"
                                   f"Apakah Anda ingin membuka file tersebut sekarang?"):
                try:
                    os.startfile(filepath)
                except Exception as e:
                    messagebox.showerror("Error", f"Gagal membuka file:\n{e}")
        except Exception as e:
            messagebox.showerror("Error", str(e))
