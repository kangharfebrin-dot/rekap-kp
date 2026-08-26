"""
Aplikasi Rekap Kegiatan Kerja Praktek
Untuk menyusun laporan kerja praktek
"""

import customtkinter as ctk
from tkinter import ttk, messagebox, filedialog
import json, os, datetime, sqlite3, hashlib
from tkcalendar import DateEntry

# === DATA MANAGER ===
DATA_DIR = os.path.dirname(os.path.abspath(__file__))
USERS_FILE = os.path.join(DATA_DIR, "users.json")
DB_FILE = os.path.join(DATA_DIR, "database_kp.db")

def init_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL
        )
    ''')
    
    cursor.execute("PRAGMA table_info(users)")
    columns = [col[1] for col in cursor.fetchall()]
    if "role" not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'user'")
        
    cursor.execute("SELECT COUNT(*) FROM users WHERE username = 'admin'")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
                       ('admin', hash_password('admin123'), 'admin'))
    
    conn.commit()
    
    if os.path.exists(USERS_FILE):
        cursor.execute("SELECT COUNT(*) FROM users")
        if cursor.fetchone()[0] <= 1:
            with open(USERS_FILE, "r", encoding="utf-8") as f:
                try:
                    users = json.load(f)
                    for user, pwd in users.items():
                        if user != 'admin':
                            cursor.execute("INSERT OR IGNORE INTO users (username, password, role) VALUES (?, ?, ?)", 
                                           (user, hash_password(pwd), 'user'))
                    conn.commit()
                    os.rename(USERS_FILE, USERS_FILE + ".bak")
                except Exception:
                    pass
    conn.close()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_login_db(username, password):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT password, role FROM users WHERE username = ?", (username,))
    row = cursor.fetchone()
    conn.close()
    if row and row[0] == hash_password(password):
        return True, row[1]
    return False, None

def register_user_db(username, password, role='user'):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("INSERT INTO users (username, password, role) VALUES (?, ?, ?)", 
                       (username, hash_password(password), role))
        conn.commit()
        success = True
    except sqlite3.IntegrityError:
        success = False
    conn.close()
    return success

def fetch_all_users_db():
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT username, role FROM users")
    rows = cursor.fetchall()
    conn.close()
    
    detailed_rows = []
    for uname, role in rows:
        nama = "-"
        nim = "-"
        perusahaan = "-"
        jml_keg = 0
        file_path = get_data_file(uname)
        if os.path.exists(file_path):
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    nama = data.get("profil", {}).get("nama", "-")
                    nim = data.get("profil", {}).get("nim", "-")
                    perusahaan = data.get("profil", {}).get("perusahaan", "-")
                    jml_keg = len(data.get("kegiatan", []))
            except Exception:
                pass
        detailed_rows.append((uname, role, nama, nim, perusahaan, jml_keg))
    return detailed_rows

def update_user_db(username, password, role):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    if password:
        cursor.execute("UPDATE users SET password = ?, role = ? WHERE username = ?", 
                       (hash_password(password), role, username))
    else:
        cursor.execute("UPDATE users SET role = ? WHERE username = ?", 
                       (role, username))
    conn.commit()
    conn.close()

def delete_user_db(username):
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE username = ?", (username,))
    conn.commit()
    conn.close()

def get_data_file(username):
    return os.path.join(DATA_DIR, f"data_kp_{username}.json")

def load_data(username):
    file_path = get_data_file(username)
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"profil": {"nama":"","nim":"","prodi":"","fakultas":"","universitas":"",
                       "perusahaan":"","alamat_perusahaan":"","periode_mulai":"","periode_selesai":"",
                       "dosen_pembimbing":"","pembimbing_lapangan":""},
            "kegiatan": []}

def save_data(username, data):
    file_path = get_data_file(username)
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# === MAIN APP ===
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("📋 Rekap Kegiatan Kerja Praktek")
        self.geometry("1100x700")
        self.minsize(900, 600)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        self.current_user = None
        self.data = None
        self.editing_idx = None
        self.admin_editing_user = None

        init_db()

        self.login_container = ctk.CTkFrame(self, fg_color="transparent")
        self.dashboard_container = ctk.CTkFrame(self, fg_color="transparent")
        self.admin_container = ctk.CTkFrame(self, fg_color="transparent")
        
        self.build_login_ui()
        self.build_dashboard_ui()
        self.build_admin_dashboard_ui()
        
        self.show_login_page()

    def save_current_data(self):
        if self.current_user and self.data:
            save_data(self.current_user, self.data)

    def show_login_page(self):
        self.dashboard_container.pack_forget()
        self.admin_container.pack_forget()
        self.login_container.pack(fill="both", expand=True)

    def show_dashboard_page(self):
        self.login_container.pack_forget()
        self.admin_container.pack_forget()
        self.dashboard_container.pack(fill="both", expand=True)
        self.data = load_data(self.current_user)
        self.refresh_ui_with_data()
        self.show_page("kegiatan")

    def show_admin_page(self):
        self.login_container.pack_forget()
        self.dashboard_container.pack_forget()
        self.admin_container.pack(fill="both", expand=True)
        self.refresh_admin_table()

    def build_login_ui(self):
        frame = ctk.CTkFrame(self.login_container, width=400, corner_radius=15)
        frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ctk.CTkLabel(frame, text="Login Akun", font=ctk.CTkFont(size=24, weight="bold")).pack(pady=(30, 20))
        
        self.user_entry = ctk.CTkEntry(frame, width=250, placeholder_text="Username")
        self.user_entry.pack(pady=10, padx=40)
        
        self.pass_entry = ctk.CTkEntry(frame, width=250, placeholder_text="Password", show="*")
        self.pass_entry.pack(pady=10, padx=40)
        
        ctk.CTkButton(frame, text="Login", width=250, font=ctk.CTkFont(weight="bold"), 
                      command=self.handle_login).pack(pady=(20, 10))
                      
        ctk.CTkButton(frame, text="Register Baru", width=250, fg_color="transparent", 
                      border_width=1, command=self.handle_register).pack(pady=(0, 30))

    def handle_login(self):
        username = self.user_entry.get().strip().lower()
        password = self.pass_entry.get()
        if not username or not password:
            messagebox.showwarning("Peringatan", "Username dan Password tidak boleh kosong!")
            return
        
        success, role = verify_login_db(username, password)
        if success:
            self.current_user = username
            self.user_entry.delete(0, "end")
            self.pass_entry.delete(0, "end")
            if role == 'admin':
                self.show_admin_page()
            else:
                self.show_dashboard_page()
        else:
            messagebox.showerror("Error", "Username atau Password salah!")

    def handle_register(self):
        username = self.user_entry.get().strip().lower()
        password = self.pass_entry.get()
        if not username or not password:
            messagebox.showwarning("Peringatan", "Username dan Password tidak boleh kosong!")
            return
        if register_user_db(username, password):
            messagebox.showinfo("Sukses", "Akun berhasil dibuat! Silakan login.")
        else:
            messagebox.showwarning("Peringatan", "Username sudah terdaftar!")

    def handle_logout(self):
        self.current_user = None
        self.data = None
        self.show_login_page()

    def build_admin_dashboard_ui(self):
        sidebar = ctk.CTkFrame(self.admin_container, width=220, corner_radius=0, fg_color="#1a1a2e")
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)

        ctk.CTkLabel(sidebar, text="👑 Admin Panel", font=ctk.CTkFont(size=22, weight="bold"),
                     text_color="#e94560").pack(pady=(30,5))
        ctk.CTkLabel(sidebar, text="Manajemen Sistem", font=ctk.CTkFont(size=12),
                     text_color="#888").pack(pady=(0,30))

        btn = ctk.CTkButton(sidebar, text="👥  Manage Users", font=ctk.CTkFont(size=15), height=45,
                            fg_color="transparent", hover_color="#e94560", anchor="w")
        btn.pack(fill="x", padx=15, pady=4)

        ctk.CTkButton(sidebar, text="🚪 Logout", font=ctk.CTkFont(size=15), height=45,
                      fg_color="transparent", hover_color="#c0392b", anchor="w",
                      command=self.handle_logout).pack(side="bottom", fill="x", padx=15, pady=20)

        main = ctk.CTkFrame(self.admin_container, fg_color="#16213e", corner_radius=0)
        main.pack(side="right", fill="both", expand=True)
        
        page = ctk.CTkFrame(main, fg_color="transparent")
        page.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(page, text="Manajemen Pengguna (CRUD)", font=ctk.CTkFont(size=20, weight="bold")).pack(anchor="w")

        # Form CRUD Admin
        admin_form = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        admin_form.pack(fill="x", pady=10)
        
        row1 = ctk.CTkFrame(admin_form, fg_color="transparent")
        row1.pack(fill="x", padx=15, pady=(12,5))
        
        ctk.CTkLabel(row1, text="Username:").pack(side="left")
        self.admin_user_entry = ctk.CTkEntry(row1, width=150)
        self.admin_user_entry.pack(side="left", padx=5)
        
        ctk.CTkLabel(row1, text="Password:").pack(side="left", padx=(10,0))
        self.admin_pass_entry = ctk.CTkEntry(row1, width=150, show="*")
        self.admin_pass_entry.pack(side="left", padx=5)
        
        ctk.CTkLabel(row1, text="Role:").pack(side="left", padx=(10,0))
        self.admin_role_var = ctk.StringVar(value="user")
        ctk.CTkOptionMenu(row1, variable=self.admin_role_var, values=["user", "admin"], width=100, fg_color="#e94560").pack(side="left", padx=5)
        
        btn_row = ctk.CTkFrame(admin_form, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=(5,12))
        
        self.admin_add_btn = ctk.CTkButton(btn_row, text="➕ Tambah Akun", fg_color="#e94560", hover_color="#c81e45", command=self.admin_add_user)
        self.admin_add_btn.pack(side="left", padx=(0,8))
        
        ctk.CTkButton(btn_row, text="🗑️ Hapus Akun", fg_color="#c0392b", hover_color="#922b21", command=self.admin_delete_user).pack(side="left", padx=(0,8))
        ctk.CTkButton(btn_row, text="🔄 Batal Edit", fg_color="#555", hover_color="#777", command=self.admin_cancel_edit).pack(side="left")

        # Table
        table_frame = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        table_frame.pack(fill="both", expand=True, pady=15)

        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview", background="#1a1a2e", foreground="white",
                        fieldbackground="#1a1a2e", rowheight=30, font=("Segoe UI", 10))
        style.configure("Treeview.Heading", background="#e94560", foreground="white",
                        font=("Segoe UI", 10, "bold"))
        style.map("Treeview", background=[("selected", "#e94560")])

        cols = ("no", "username", "role", "nama", "nim", "perusahaan", "jml")
        self.admin_tree = ttk.Treeview(table_frame, columns=cols, show="headings", selectmode="browse")

        for cid, txt, w in [("no","No",40), ("username","Username",120), ("role","Role",80),
                            ("nama","Nama Lengkap",150), ("nim","NIM",90),
                            ("perusahaan","Perusahaan",150), ("jml","Jml Keg",70)]:
            self.admin_tree.heading(cid, text=txt)
            self.admin_tree.column(cid, width=w, minwidth=50)

        sb = ttk.Scrollbar(table_frame, orient="vertical", command=self.admin_tree.yview)
        self.admin_tree.configure(yscrollcommand=sb.set)
        self.admin_tree.pack(side="left", fill="both", expand=True, padx=(10,0), pady=10)
        sb.pack(side="right", fill="y", pady=10, padx=(0,10))
        
        self.admin_tree.bind("<Double-1>", self.on_admin_double_click)
                      
        self.admin_count_label = ctk.CTkLabel(page, text="", font=ctk.CTkFont(size=12), text_color="#888")
        self.admin_count_label.pack(anchor="w", pady=5)

    def on_admin_double_click(self, event):
        sel = self.admin_tree.selection()
        if not sel: return
        item = self.admin_tree.item(sel[0])
        uname = item['values'][1]
        role = item['values'][2]
        
        self.admin_editing_user = uname
        self.admin_add_btn.configure(text="✏️ Update Akun")
        self.admin_user_entry.delete(0, "end")
        self.admin_user_entry.insert(0, uname)
        self.admin_user_entry.configure(state="disabled")
        self.admin_pass_entry.delete(0, "end")
        self.admin_role_var.set(role)
        
    def admin_cancel_edit(self):
        self.admin_editing_user = None
        self.admin_add_btn.configure(text="➕ Tambah Akun")
        self.admin_user_entry.configure(state="normal")
        self.admin_user_entry.delete(0, "end")
        self.admin_pass_entry.delete(0, "end")
        self.admin_role_var.set("user")
        
    def admin_add_user(self):
        uname = self.admin_user_entry.get().strip().lower()
        pwd = self.admin_pass_entry.get()
        role = self.admin_role_var.get()
        
        if not uname:
            messagebox.showwarning("Peringatan", "Username tidak boleh kosong!")
            return
            
        if self.admin_editing_user:
            if self.admin_editing_user == self.current_user and role != 'admin':
                messagebox.showwarning("Peringatan", "Anda tidak bisa mengubah role Anda sendiri menjadi user!")
                return
            update_user_db(self.admin_editing_user, pwd, role)
            messagebox.showinfo("Sukses", f"Akun '{self.admin_editing_user}' berhasil diupdate!")
            self.admin_cancel_edit()
            self.refresh_admin_table()
        else:
            if not pwd:
                messagebox.showwarning("Peringatan", "Password tidak boleh kosong untuk akun baru!")
                return
            if register_user_db(uname, pwd, role):
                messagebox.showinfo("Sukses", f"Akun '{uname}' berhasil dibuat!")
                self.admin_cancel_edit()
                self.refresh_admin_table()
            else:
                messagebox.showwarning("Peringatan", "Username sudah terdaftar!")

    def refresh_admin_table(self):
        for item in self.admin_tree.get_children():
            self.admin_tree.delete(item)
        users = fetch_all_users_db()
        for i, u in enumerate(users):
            self.admin_tree.insert("", "end", values=(i+1, u[0], u[1], u[2], u[3], u[4], u[5]))
        self.admin_count_label.configure(text=f"Total: {len(users)} pengguna terdaftar")

    def admin_delete_user(self):
        sel = self.admin_tree.selection()
        if not sel:
            messagebox.showwarning("Peringatan", "Pilih user yang ingin dihapus!")
            return
        idx = self.admin_tree.index(sel[0])
        item = self.admin_tree.item(sel[0])
        uname = item['values'][1]
        role = item['values'][2]
        
        if role == 'admin' and uname == self.current_user:
            messagebox.showwarning("Peringatan", "Anda tidak bisa menghapus akun Anda sendiri saat sedang login!")
            return
            
        if messagebox.askyesno("Konfirmasi", f"Apakah Anda yakin ingin menghapus pengguna '{uname}'?\nCatatan: Data laporan mereka akan tetap tersimpan sebagai arsip."):
            delete_user_db(uname)
            self.admin_cancel_edit()
            self.refresh_admin_table()
            messagebox.showinfo("Sukses", f"Pengguna '{uname}' berhasil dihapus.")

    def build_dashboard_ui(self):
        sidebar = ctk.CTkFrame(self.dashboard_container, width=220, corner_radius=0, fg_color="#1a1a2e")
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)

        ctk.CTkLabel(sidebar, text="📋 Rekap KP", font=ctk.CTkFont(size=22, weight="bold"),
                     text_color="#e94560").pack(pady=(30,5))
        ctk.CTkLabel(sidebar, text="Kerja Praktek Logger", font=ctk.CTkFont(size=12),
                     text_color="#888").pack(pady=(0,30))

        self.main_frame = ctk.CTkFrame(self.dashboard_container, fg_color="#16213e", corner_radius=0)
        self.main_frame.pack(side="right", fill="both", expand=True)

        self.pages = {}
        btn_data = [("👤  Profil", "profil"), ("📝  Kegiatan", "kegiatan"), ("📤  Export", "export")]
        for text, name in btn_data:
            btn = ctk.CTkButton(sidebar, text=text, font=ctk.CTkFont(size=15), height=45,
                                fg_color="transparent", hover_color="#e94560", anchor="w",
                                command=lambda n=name: self.show_page(n))
            btn.pack(fill="x", padx=15, pady=4)

        ctk.CTkButton(sidebar, text="🚪 Logout", font=ctk.CTkFont(size=15), height=45,
                      fg_color="transparent", hover_color="#c0392b", anchor="w",
                      command=self.handle_logout).pack(side="bottom", fill="x", padx=15, pady=20)

        self.build_profil_page()
        self.build_kegiatan_page()
        self.build_export_page()

    def refresh_ui_with_data(self):
        for key, entry in self.profil_entries.items():
            entry.delete(0, "end")
            entry.insert(0, self.data["profil"].get(key, ""))
        self.refresh_table()

    def show_page(self, name):
        for p in self.pages.values():
            p.pack_forget()
        self.pages[name].pack(fill="both", expand=True, padx=20, pady=20)

    # === PROFIL PAGE ===
    def build_profil_page(self):
        page = ctk.CTkScrollableFrame(self.main_frame, fg_color="transparent")
        self.pages["profil"] = page

        ctk.CTkLabel(page, text="Data Profil Kerja Praktek", font=ctk.CTkFont(size=20, weight="bold")).pack(anchor="w", pady=(0,15))

        self.profil_entries = {}
        sections = [
            ("Data Mahasiswa", [("Nama Lengkap","nama"),("NIM","nim"),("Program Studi","prodi"),
                                ("Fakultas","fakultas"),("Universitas","universitas")]),
            ("Data Perusahaan/Instansi", [("Nama Perusahaan","perusahaan"),("Alamat","alamat_perusahaan"),
                                          ("Periode Mulai (dd/mm/yyyy)","periode_mulai"),("Periode Selesai (dd/mm/yyyy)","periode_selesai")]),
            ("Data Pembimbing", [("Dosen Pembimbing","dosen_pembimbing"),("Pembimbing Lapangan","pembimbing_lapangan")])
        ]

        for sec_title, fields in sections:
            f = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
            f.pack(fill="x", pady=8)
            ctk.CTkLabel(f, text=sec_title, font=ctk.CTkFont(size=15, weight="bold"),
                         text_color="#e94560").pack(anchor="w", padx=15, pady=(12,5))
            for label, key in fields:
                row = ctk.CTkFrame(f, fg_color="transparent")
                row.pack(fill="x", padx=15, pady=3)
                ctk.CTkLabel(row, text=label, width=220, anchor="w").pack(side="left")
                e = ctk.CTkEntry(row, width=400, placeholder_text=f"Masukkan {label.lower()}")
                e.pack(side="left", padx=(10,0))
                self.profil_entries[key] = e

        ctk.CTkButton(page, text="💾  Simpan Profil", font=ctk.CTkFont(size=14, weight="bold"),
                      height=42, fg_color="#e94560", hover_color="#c81e45",
                      command=self.save_profil).pack(pady=15)

    def save_profil(self):
        for key, entry in self.profil_entries.items():
            self.data["profil"][key] = entry.get().strip()
        self.save_current_data()
        messagebox.showinfo("Sukses", "Profil berhasil disimpan!")

    # === KEGIATAN PAGE ===
    def build_kegiatan_page(self):
        page = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.pages["kegiatan"] = page

        ctk.CTkLabel(page, text="Catatan Kegiatan Harian", font=ctk.CTkFont(size=20, weight="bold")).pack(anchor="w")

        # Form
        form = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        form.pack(fill="x", pady=10)

        row1 = ctk.CTkFrame(form, fg_color="transparent")
        row1.pack(fill="x", padx=15, pady=(12,5))

        ctk.CTkLabel(row1, text="Tanggal:").pack(side="left")
        self.date_entry = DateEntry(row1, width=14, date_pattern="dd/mm/yyyy",
                                    background="#e94560", foreground="white")
        self.date_entry.pack(side="left", padx=(5,20))

        ctk.CTkLabel(row1, text="Minggu ke:").pack(side="left")
        self.minggu_var = ctk.StringVar(value="1")
        ctk.CTkOptionMenu(row1, variable=self.minggu_var, values=[str(i) for i in range(1,25)],
                          width=80, fg_color="#e94560").pack(side="left", padx=5)

        row2 = ctk.CTkFrame(form, fg_color="transparent")
        row2.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(row2, text="Kegiatan:").pack(anchor="w")
        self.kegiatan_text = ctk.CTkTextbox(row2, height=60, fg_color="#16213e", corner_radius=8)
        self.kegiatan_text.pack(fill="x", pady=3)

        row3 = ctk.CTkFrame(form, fg_color="transparent")
        row3.pack(fill="x", padx=15, pady=5)
        ctk.CTkLabel(row3, text="Keterangan/Hasil:").pack(anchor="w")
        self.keterangan_text = ctk.CTkTextbox(row3, height=40, fg_color="#16213e", corner_radius=8)
        self.keterangan_text.pack(fill="x", pady=3)

        btn_row = ctk.CTkFrame(form, fg_color="transparent")
        btn_row.pack(fill="x", padx=15, pady=(5,12))

        self.add_btn = ctk.CTkButton(btn_row, text="➕ Tambah", fg_color="#e94560",
                                     hover_color="#c81e45", command=self.add_kegiatan)
        self.add_btn.pack(side="left", padx=(0,8))
        ctk.CTkButton(btn_row, text="🗑️ Hapus Pilihan", fg_color="#c0392b",
                      hover_color="#922b21", command=self.delete_kegiatan).pack(side="left", padx=(0,8))
        ctk.CTkButton(btn_row, text="🔄 Batal Edit", fg_color="#555",
                      hover_color="#777", command=self.cancel_edit).pack(side="left")

        # Table
        table_frame = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        table_frame.pack(fill="both", expand=True, pady=5)

        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview", background="#1a1a2e", foreground="white",
                        fieldbackground="#1a1a2e", rowheight=30, font=("Segoe UI", 10))
        style.configure("Treeview.Heading", background="#e94560", foreground="white",
                        font=("Segoe UI", 10, "bold"))
        style.map("Treeview", background=[("selected", "#e94560")])

        cols = ("no","tanggal","minggu","kegiatan","keterangan")
        self.tree = ttk.Treeview(table_frame, columns=cols, show="headings", selectmode="browse")

        for cid, txt, w in [("no","No",40),("tanggal","Tanggal",100),("minggu","Minggu",70),
                             ("kegiatan","Kegiatan",350),("keterangan","Keterangan",250)]:
            self.tree.heading(cid, text=txt)
            self.tree.column(cid, width=w, minwidth=40)

        sb = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side="left", fill="both", expand=True, padx=(10,0), pady=10)
        sb.pack(side="right", fill="y", pady=10, padx=(0,10))

        self.tree.bind("<Double-1>", self.on_double_click)

        # Count label
        self.count_label = ctk.CTkLabel(page, text="", font=ctk.CTkFont(size=12), text_color="#888")
        self.count_label.pack(anchor="w")

    def refresh_table(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        if self.data and "kegiatan" in self.data:
            for i, k in enumerate(self.data["kegiatan"]):
                self.tree.insert("", "end", values=(i+1, k["tanggal"], f"Minggu {k['minggu']}",
                                                     k["kegiatan"], k["keterangan"]))
        self.update_count()

    def update_count(self):
        if hasattr(self, "count_label") and self.data and "kegiatan" in self.data:
            self.count_label.configure(text=f"Total: {len(self.data['kegiatan'])} kegiatan tercatat")

    def add_kegiatan(self):
        tanggal = self.date_entry.get()
        minggu = self.minggu_var.get()
        kegiatan = self.kegiatan_text.get("1.0", "end").strip()
        keterangan = self.keterangan_text.get("1.0", "end").strip()

        if not kegiatan:
            messagebox.showwarning("Peringatan", "Kegiatan tidak boleh kosong!")
            return

        entry = {"tanggal": tanggal, "minggu": minggu, "kegiatan": kegiatan, "keterangan": keterangan}

        if self.editing_idx is not None:
            self.data["kegiatan"][self.editing_idx] = entry
            self.editing_idx = None
            self.add_btn.configure(text="➕ Tambah")
        else:
            self.data["kegiatan"].append(entry)

        self.save_current_data()
        self.clear_form()
        self.refresh_table()

    def delete_kegiatan(self):
        sel = self.tree.selection()
        if not sel:
            messagebox.showwarning("Peringatan", "Pilih kegiatan yang ingin dihapus!")
            return
        idx = self.tree.index(sel[0])
        if messagebox.askyesno("Konfirmasi", "Hapus kegiatan ini?"):
            self.data["kegiatan"].pop(idx)
            self.save_current_data()
            self.cancel_edit()
            self.refresh_table()

    def on_double_click(self, event):
        sel = self.tree.selection()
        if not sel:
            return
        idx = self.tree.index(sel[0])
        k = self.data["kegiatan"][idx]
        self.editing_idx = idx
        self.add_btn.configure(text="✏️ Update")

        self.date_entry.delete(0, "end")
        self.date_entry.insert(0, k["tanggal"])
        self.minggu_var.set(k["minggu"])
        self.kegiatan_text.delete("1.0", "end")
        self.kegiatan_text.insert("1.0", k["kegiatan"])
        self.keterangan_text.delete("1.0", "end")
        self.keterangan_text.insert("1.0", k["keterangan"])

    def cancel_edit(self):
        self.editing_idx = None
        self.add_btn.configure(text="➕ Tambah")
        self.clear_form()

    def clear_form(self):
        self.kegiatan_text.delete("1.0", "end")
        self.keterangan_text.delete("1.0", "end")

    # === EXPORT PAGE ===
    def build_export_page(self):
        page = ctk.CTkScrollableFrame(self.main_frame, fg_color="transparent")
        self.pages["export"] = page

        ctk.CTkLabel(page, text="Export Laporan", font=ctk.CTkFont(size=20, weight="bold")).pack(anchor="w", pady=(0,15))

        # Stats
        stats = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        stats.pack(fill="x", pady=8)
        ctk.CTkLabel(stats, text="📊 Ringkasan Data", font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e94560").pack(anchor="w", padx=15, pady=(12,5))
        self.stats_label = ctk.CTkLabel(stats, text="", justify="left", font=ctk.CTkFont(size=13))
        self.stats_label.pack(anchor="w", padx=15, pady=(0,12))

        # Export buttons
        exp = ctk.CTkFrame(page, fg_color="#1a1a2e", corner_radius=12)
        exp.pack(fill="x", pady=8)
        ctk.CTkLabel(exp, text="📤 Pilih Format Export", font=ctk.CTkFont(size=15, weight="bold"),
                     text_color="#e94560").pack(anchor="w", padx=15, pady=(12,5))

        for txt, color, hover, cmd in [
            ("📊  Export ke Excel (.xlsx)", "#27ae60", "#1e8449", self.export_excel),
            ("📄  Export ke Word (.docx)", "#2980b9", "#1f6fa3", self.export_word)
        ]:
            ctk.CTkButton(exp, text=txt, font=ctk.CTkFont(size=14, weight="bold"), height=50,
                          fg_color=color, hover_color=hover, command=cmd).pack(fill="x", padx=15, pady=6)

        ctk.CTkLabel(exp, text="", height=10).pack()  # spacer

        # Update stats when shown
        page.bind("<Map>", lambda e: self.update_stats())

    def update_stats(self):
        p = self.data["profil"]
        n = len(self.data["kegiatan"])
        weeks = set(k["minggu"] for k in self.data["kegiatan"])
        txt = (f"Nama: {p['nama'] or '-'}  |  NIM: {p['nim'] or '-'}\n"
               f"Perusahaan: {p['perusahaan'] or '-'}\n"
               f"Total Kegiatan: {n}  |  Jumlah Minggu: {len(weeks)}")
        self.stats_label.configure(text=txt)

    def export_excel(self):
        if not self.data["kegiatan"]:
            messagebox.showwarning("Peringatan", "Belum ada data kegiatan!")
            return
        filepath = filedialog.asksaveasfilename(defaultextension=".xlsx",
                    filetypes=[("Excel", "*.xlsx")], initialfile="Rekap_KP.xlsx")
        if not filepath:
            return
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

            wb = Workbook()
            ws = wb.active
            ws.title = "Rekap Kegiatan KP"

            # Header info
            p = self.data["profil"]
            ws.merge_cells("A1:E1")
            ws["A1"] = "REKAP KEGIATAN KERJA PRAKTEK"
            ws["A1"].font = Font(size=14, bold=True)
            ws["A1"].alignment = Alignment(horizontal="center")

            info = [f"Nama: {p['nama']}", f"NIM: {p['nim']}", f"Prodi: {p['prodi']}",
                    f"Perusahaan: {p['perusahaan']}", f"Periode: {p['periode_mulai']} - {p['periode_selesai']}"]
            for i, txt in enumerate(info):
                ws.merge_cells(f"A{i+2}:E{i+2}")
                ws[f"A{i+2}"] = txt

            # Table header
            start_row = len(info) + 3
            headers = ["No", "Tanggal", "Minggu", "Kegiatan", "Keterangan"]
            header_fill = PatternFill(start_color="E94560", end_color="E94560", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            thin = Side(style="thin")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            for j, h in enumerate(headers):
                cell = ws.cell(row=start_row, column=j+1, value=h)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                cell.border = border

            for i, k in enumerate(self.data["kegiatan"]):
                row = start_row + 1 + i
                vals = [i+1, k["tanggal"], f"Minggu {k['minggu']}", k["kegiatan"], k["keterangan"]]
                for j, v in enumerate(vals):
                    cell = ws.cell(row=row, column=j+1, value=v)
                    cell.border = border
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            ws.column_dimensions["A"].width = 5
            ws.column_dimensions["B"].width = 14
            ws.column_dimensions["C"].width = 12
            ws.column_dimensions["D"].width = 45
            ws.column_dimensions["E"].width = 30

            wb.save(filepath)
            if messagebox.askyesno("Sukses", f"File Excel berhasil disimpan!\n{filepath}\n\nApakah Anda ingin membuka file tersebut sekarang?"):
                try:
                    os.startfile(filepath)
                except Exception as e:
                    messagebox.showerror("Error", f"Gagal membuka file:\n{e}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def export_word(self):
        if not self.data["kegiatan"]:
            messagebox.showwarning("Peringatan", "Belum ada data kegiatan!")
            return
        filepath = filedialog.asksaveasfilename(defaultextension=".docx",
                    filetypes=[("Word", "*.docx")], initialfile="Laporan_KP.docx")
        if not filepath:
            return
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            p = self.data["profil"]

            # Title
            title = doc.add_heading("LAPORAN KEGIATAN KERJA PRAKTEK", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Profile section
            doc.add_heading("Data Mahasiswa", level=2)
            for label, key in [("Nama", "nama"), ("NIM", "nim"), ("Program Studi", "prodi"),
                               ("Fakultas", "fakultas"), ("Universitas", "universitas")]:
                doc.add_paragraph(f"{label}: {p.get(key, '-')}")

            doc.add_heading("Data Tempat Kerja Praktek", level=2)
            for label, key in [("Perusahaan/Instansi", "perusahaan"), ("Alamat", "alamat_perusahaan"),
                               ("Periode", None), ("Dosen Pembimbing", "dosen_pembimbing"),
                               ("Pembimbing Lapangan", "pembimbing_lapangan")]:
                if key:
                    doc.add_paragraph(f"{label}: {p.get(key, '-')}")
                else:
                    doc.add_paragraph(f"Periode: {p.get('periode_mulai','-')} s/d {p.get('periode_selesai','-')}")

            # Activities grouped by week
            doc.add_heading("Catatan Kegiatan Harian", level=2)

            weeks = {}
            for k in self.data["kegiatan"]:
                w = k["minggu"]
                weeks.setdefault(w, []).append(k)

            for w in sorted(weeks.keys(), key=lambda x: int(x)):
                doc.add_heading(f"Minggu ke-{w}", level=3)
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, txt in enumerate(["No", "Tanggal", "Kegiatan", "Keterangan"]):
                    cell = table.rows[0].cells[i]
                    cell.text = txt
                    for par in cell.paragraphs:
                        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in par.runs:
                            run.bold = True

                for j, k in enumerate(weeks[w]):
                    row = table.add_row()
                    row.cells[0].text = str(j+1)
                    row.cells[1].text = k["tanggal"]
                    row.cells[2].text = k["kegiatan"]
                    row.cells[3].text = k["keterangan"]

                doc.add_paragraph("")  # spacer

            doc.save(filepath)
            if messagebox.askyesno("Sukses", f"File Word berhasil disimpan!\n{filepath}\n\nApakah Anda ingin membuka file tersebut sekarang?"):
                try:
                    os.startfile(filepath)
                except Exception as e:
                    messagebox.showerror("Error", f"Gagal membuka file:\n{e}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    app = App()
    app.mainloop()
